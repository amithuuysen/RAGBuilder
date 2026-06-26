import logging
import os
import re
import math
from typing import Any, Dict, List, Optional, Tuple
import pypdf
import docx
from langchain_core.documents import Document
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_qdrant import Qdrant
from langchain_community.retrievers import BM25Retriever
from langchain.retrievers import EnsembleRetriever
from qdrant_client import QdrantClient
from qdrant_client.http import models as qdrant_models

from lib.config import QDRANT_PATH
from lib.metadata import build_qdrant_filter, compute_facets

logger = logging.getLogger(__name__)

_client: Optional[QdrantClient] = None
_bm25_cache: Dict[int, Tuple[List[Document], BM25Retriever]] = {}


def get_qdrant_client() -> QdrantClient:
    global _client
    if _client is None:
        os.makedirs(QDRANT_PATH, exist_ok=True)
        _client = QdrantClient(path=QDRANT_PATH)
    return _client


def invalidate_bm25_cache(bot_id: int):
    _bm25_cache.pop(bot_id, None)


def _source_type_from_filename(filename: str) -> str:
    ext = os.path.splitext(filename.lower())[1]
    mapping = {".pdf": "pdf", ".docx": "docx", ".doc": "docx", ".txt": "txt", ".md": "md"}
    return mapping.get(ext, "unknown")


def parse_file(filepath: str, filename: str) -> Tuple[str, List[Dict[str, Any]]]:
    """Extract text and optional page segments from supported file types."""
    ext = os.path.splitext(filename.lower())[1]
    segments: List[Dict[str, Any]] = []

    if ext == ".pdf":
        text_parts = []
        try:
            reader = pypdf.PdfReader(filepath)
            for page_num, page in enumerate(reader.pages, start=1):
                extracted = page.extract_text() or ""
                if extracted.strip():
                    text_parts.append(extracted)
                    segments.append({
                        "text": extracted,
                        "page_number": page_num,
                        "section_heading": "",
                    })
        except Exception as e:
            raise ValueError(f"Failed to parse PDF file: {e}")
        return "\n".join(text_parts), segments

    if ext in (".docx", ".doc"):
        try:
            doc = docx.Document(filepath)
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            full_text = "\n".join(paragraphs)
            current_heading = ""
            current_lines: List[str] = []
            for para in doc.paragraphs:
                if not para.text.strip():
                    continue
                style_name = (para.style.name or "").lower() if para.style else ""
                if "heading" in style_name:
                    if current_lines:
                        segments.append({
                            "text": "\n".join(current_lines),
                            "page_number": None,
                            "section_heading": current_heading,
                        })
                        current_lines = []
                    current_heading = para.text.strip()
                current_lines.append(para.text)
            if current_lines:
                segments.append({
                    "text": "\n".join(current_lines),
                    "page_number": None,
                    "section_heading": current_heading,
                })
            if not segments:
                segments.append({"text": full_text, "page_number": None, "section_heading": ""})
            return full_text, segments
        except Exception as e:
            raise ValueError(f"Failed to parse DOCX file: {e}")

    try:
        with open(filepath, "r", encoding="utf-8", errors="ignore") as f:
            content = f.read()
    except Exception as e:
        raise ValueError(f"Failed to read text file: {e}")

    if ext == ".md":
        parts = re.split(r"(?=^#{1,3}\s)", content, flags=re.MULTILINE)
        for part in parts:
            part = part.strip()
            if not part:
                continue
            heading_match = re.match(r"^(#{1,3})\s+(.+)$", part.split("\n")[0])
            heading = heading_match.group(2).strip() if heading_match else ""
            segments.append({"text": part, "page_number": None, "section_heading": heading})
    else:
        segments.append({"text": content, "page_number": None, "section_heading": ""})

    return content, segments


def chunk_document(
    text: str,
    chunk_size: int = 500,
    chunk_overlap: int = 50,
    segments: Optional[List[Dict[str, Any]]] = None,
) -> List[Document]:
    splitter = RecursiveCharacterTextSplitter(
        chunk_size=chunk_size,
        chunk_overlap=chunk_overlap,
        length_function=len,
        separators=["\n\n", "\n", " ", ""],
    )

    if not segments:
        return splitter.create_documents([text])

    docs: List[Document] = []
    for segment in segments:
        segment_text = segment.get("text", "")
        if not segment_text.strip():
            continue
        for doc in splitter.create_documents([segment_text]):
            doc.metadata["page_number"] = segment.get("page_number")
            doc.metadata["section_heading"] = segment.get("section_heading") or ""
            docs.append(doc)
    return docs


def index_document(
    bot_id: int,
    doc_id: int,
    text: str,
    name: str,
    embeddings_model,
    chunk_size: int = 500,
    chunk_overlap: int = 50,
    doc_metadata: Optional[Dict[str, Any]] = None,
    segments: Optional[List[Dict[str, Any]]] = None,
) -> int:
    client = get_qdrant_client()
    collection_name = f"bot_{bot_id}"
    base_meta = doc_metadata or {}

    docs = chunk_document(text, chunk_size, chunk_overlap, segments)
    for i, doc in enumerate(docs):
        doc.metadata.update({
            "doc_id": doc_id,
            "bot_id": bot_id,
            "doc_name": name,
            "chunk_index": i,
            "source_type": base_meta.get("source_type", _source_type_from_filename(name)),
            "tags": base_meta.get("tags", []),
            "source_url": base_meta.get("source_url", ""),
            "author": base_meta.get("author", ""),
            "uploaded_at": base_meta.get("uploaded_at", ""),
            "custom_fields": base_meta.get("custom_fields", {}),
        })
        if doc.metadata.get("page_number") is None:
            doc.metadata.pop("page_number", None)

    if not docs:
        return 0

    try:
        collections = client.get_collections().collections
        exists = any(c.name == collection_name for c in collections)
        if not exists:
            dummy_vector = embeddings_model.embed_query("test")
            client.create_collection(
                collection_name=collection_name,
                vectors_config=qdrant_models.VectorParams(
                    size=len(dummy_vector),
                    distance=qdrant_models.Distance.COSINE,
                ),
            )
    except Exception as e:
        raise RuntimeError(f"Failed to check/create Qdrant collection: {e}") from e

    vector_store = Qdrant(
        client=client,
        collection_name=collection_name,
        embeddings=embeddings_model,
    )
    vector_store.add_documents(docs)
    invalidate_bm25_cache(bot_id)
    return len(docs)


def delete_document_chunks(bot_id: int, doc_id: int):
    client = get_qdrant_client()
    collection_name = f"bot_{bot_id}"
    try:
        collections = client.get_collections().collections
        if not any(c.name == collection_name for c in collections):
            return
        client.delete(
            collection_name=collection_name,
            points_selector=qdrant_models.Filter(
                must=[
                    qdrant_models.FieldCondition(
                        key="metadata.doc_id",
                        match=qdrant_models.MatchValue(value=doc_id),
                    )
                ]
            ),
        )
        invalidate_bm25_cache(bot_id)
    except Exception as e:
        logger.warning("Error deleting chunks from Qdrant: %s", e)


def delete_bot_collection(bot_id: int):
    client = get_qdrant_client()
    collection_name = f"bot_{bot_id}"
    try:
        collections = client.get_collections().collections
        if any(c.name == collection_name for c in collections):
            client.delete_collection(collection_name)
        invalidate_bm25_cache(bot_id)
    except Exception as e:
        logger.warning("Error deleting Qdrant collection for bot %s: %s", bot_id, e)


def _scroll_all_documents(client: QdrantClient, collection_name: str) -> List[Document]:
    all_docs: List[Document] = []
    offset = None
    while True:
        response = client.scroll(
            collection_name=collection_name,
            limit=1000,
            offset=offset,
            with_payload=True,
            with_vectors=False,
        )
        points, offset = response
        for p in points:
            if "page_content" in p.payload:
                all_docs.append(
                    Document(
                        page_content=p.payload["page_content"],
                        metadata=p.payload.get("metadata", {}),
                    )
                )
        if offset is None:
            break
    return all_docs


def _get_bm25_retriever(bot_id: int, client: QdrantClient, collection_name: str) -> Optional[BM25Retriever]:
    if bot_id in _bm25_cache:
        return _bm25_cache[bot_id][1]

    all_docs = _scroll_all_documents(client, collection_name)
    if not all_docs:
        return None

    retriever = BM25Retriever.from_documents(all_docs)
    _bm25_cache[bot_id] = (all_docs, retriever)
    return retriever


def _cosine_similarity(a: List[float], b: List[float]) -> float:
    dot = sum(x * y for x, y in zip(a, b))
    norm_a = math.sqrt(sum(x * x for x in a))
    norm_b = math.sqrt(sum(x * x for x in b))
    if norm_a == 0 or norm_b == 0:
        return 0.0
    return dot / (norm_a * norm_b)


def _apply_mmr(
    query_vector: List[float],
    candidates: List[Dict[str, Any]],
    vectors: Dict[str, List[float]],
    top_k: int,
    lambda_mult: float = 0.5,
) -> List[Dict[str, Any]]:
    if not candidates:
        return []
    selected: List[Dict[str, Any]] = []
    remaining = candidates.copy()

    while remaining and len(selected) < top_k:
        best_score = -1.0
        best_idx = 0
        for idx, cand in enumerate(remaining):
            point_id = cand.get("_point_id")
            vec = vectors.get(point_id)
            if vec is None:
                relevance = cand.get("vec_score", 0.0)
            else:
                relevance = _cosine_similarity(query_vector, vec)

            if not selected:
                mmr_score = relevance
            else:
                max_sim = 0.0
                for sel in selected:
                    sel_id = sel.get("_point_id")
                    sel_vec = vectors.get(sel_id)
                    if vec is not None and sel_vec is not None:
                        max_sim = max(max_sim, _cosine_similarity(vec, sel_vec))
                mmr_score = lambda_mult * relevance - (1 - lambda_mult) * max_sim

            if mmr_score > best_score:
                best_score = mmr_score
                best_idx = idx

        selected.append(remaining.pop(best_idx))

    return selected


def _vector_search_with_scores(
    client: QdrantClient,
    collection_name: str,
    query_vector: List[float],
    top_k: int,
    qdrant_filter: Optional[qdrant_models.Filter],
) -> List[Tuple[Any, float]]:
    results = client.search(
        collection_name=collection_name,
        query_vector=query_vector,
        limit=top_k * 3 if top_k else top_k,
        query_filter=qdrant_filter,
        with_payload=True,
    )
    return [(hit, hit.score) for hit in results]


def retrieve_context(
    bot_id: int,
    query: str,
    embeddings_model,
    technique: str = "vector",
    top_k: int = 4,
    vector_weight: float = 0.7,
    keyword_weight: float = 0.3,
    score_threshold: float = 0.0,
    use_mmr: bool = False,
    filters: Optional[Dict[str, Any]] = None,
) -> List[Dict[str, Any]]:
    client = get_qdrant_client()
    collection_name = f"bot_{bot_id}"

    try:
        collections = client.get_collections().collections
        if not any(c.name == collection_name for c in collections):
            return []
    except Exception:
        return []

    qdrant_filter = build_qdrant_filter(filters)
    query_vector = embeddings_model.embed_query(query)

    vector_hits = _vector_search_with_scores(
        client, collection_name, query_vector, top_k, qdrant_filter
    )

    vec_results: Dict[str, Dict[str, Any]] = {}
    vectors_map: Dict[str, List[float]] = {}

    for hit, score in vector_hits:
        payload = hit.payload or {}
        content = payload.get("page_content", "")
        metadata = payload.get("metadata", {})
        key = f"{metadata.get('doc_id')}:{metadata.get('chunk_index')}"
        vec_results[key] = {
            "content": content,
            "metadata": metadata,
            "vec_score": float(score),
            "kw_score": 0.0,
            "_point_id": str(hit.id),
        }
        if hit.vector:
            vectors_map[str(hit.id)] = hit.vector

    kw_results: Dict[str, float] = {}
    if technique in ("keyword", "hybrid"):
        try:
            bm25 = _get_bm25_retriever(bot_id, client, collection_name)
            if bm25:
                bm25.k = top_k * 3
                kw_docs = bm25.invoke(query)
                max_rank = len(kw_docs) or 1
                for rank, doc in enumerate(kw_docs):
                    meta = doc.metadata
                    if filters and not _doc_matches_filter(meta, filters):
                        continue
                    key = f"{meta.get('doc_id')}:{meta.get('chunk_index')}"
                    kw_results[key] = 1.0 - (rank / max_rank)
        except Exception as e:
            logger.warning("BM25 retrieval failed: %s", e)

    combined: Dict[str, Dict[str, Any]] = {}

    if technique == "vector":
        for key, item in vec_results.items():
            combined[key] = item
    elif technique == "keyword":
        all_docs = _scroll_all_documents(client, collection_name)
        doc_map = {
            f"{d.metadata.get('doc_id')}:{d.metadata.get('chunk_index')}": d
            for d in all_docs
        }
        for key, kw_score in sorted(kw_results.items(), key=lambda x: -x[1])[: top_k * 3]:
            doc = doc_map.get(key)
            if doc:
                combined[key] = {
                    "content": doc.page_content,
                    "metadata": doc.metadata,
                    "vec_score": 0.0,
                    "kw_score": kw_score,
                    "_point_id": key,
                }
    else:
        all_keys = set(vec_results.keys()) | set(kw_results.keys())
        total_weight = vector_weight + keyword_weight or 1.0
        vw = vector_weight / total_weight
        kw = keyword_weight / total_weight
        all_docs = _scroll_all_documents(client, collection_name)
        doc_map = {
            f"{d.metadata.get('doc_id')}:{d.metadata.get('chunk_index')}": d
            for d in all_docs
        }
        for key in all_keys:
            vec_score = vec_results.get(key, {}).get("vec_score", 0.0)
            kw_score = kw_results.get(key, 0.0)
            combined_score = vw * vec_score + kw * kw_score
            if key in vec_results:
                item = vec_results[key].copy()
            else:
                doc = doc_map.get(key)
                if not doc:
                    continue
                item = {
                    "content": doc.page_content,
                    "metadata": doc.metadata,
                    "vec_score": 0.0,
                    "kw_score": kw_score,
                    "_point_id": key,
                }
            item["combined_score"] = combined_score
            item["kw_score"] = kw_score
            item["vec_score"] = vec_score
            combined[key] = item

    results = list(combined.values())
    if technique != "hybrid":
        for item in results:
            item["combined_score"] = item.get("vec_score", 0.0) if technique == "vector" else item.get("kw_score", 0.0)

    results.sort(key=lambda x: x.get("combined_score", 0.0), reverse=True)

    if score_threshold > 0:
        results = [r for r in results if r.get("combined_score", 0.0) >= score_threshold]

    if use_mmr and technique != "keyword":
        candidate_ids = [
            r["_point_id"] for r in results[: top_k * 3]
            if r.get("_point_id") and ":" not in str(r["_point_id"])
        ]
        if candidate_ids:
            fetch_vectors = client.retrieve(
                collection_name=collection_name,
                ids=candidate_ids,
                with_vectors=True,
            )
            for point in fetch_vectors:
                if point.vector:
                    vectors_map[str(point.id)] = point.vector
        results = _apply_mmr(query_vector, results[: top_k * 3], vectors_map, top_k)
    else:
        results = results[:top_k]

    for item in results:
        item.pop("_point_id", None)
        if "combined_score" not in item:
            item["combined_score"] = item.get("vec_score", 0.0)

    return results


def _doc_matches_filter(metadata: Dict[str, Any], filters: Dict[str, Any]) -> bool:
    if filters.get("doc_ids") and metadata.get("doc_id") not in filters["doc_ids"]:
        return False
    if filters.get("exclude_doc_ids") and metadata.get("doc_id") in filters["exclude_doc_ids"]:
        return False
    if filters.get("source_type") and metadata.get("source_type") != filters["source_type"]:
        return False
    if filters.get("author") and metadata.get("author") != filters["author"]:
        return False
    tags_filter = filters.get("tags")
    if tags_filter:
        doc_tags = metadata.get("tags") or []
        values = tags_filter.get("values", [])
        mode = tags_filter.get("mode", "any")
        if mode == "all":
            if not all(t in doc_tags for t in values):
                return False
        elif not any(t in doc_tags for t in values):
            return False
    return True


def get_collection_facets(bot_id: int) -> Dict[str, Any]:
    client = get_qdrant_client()
    collection_name = f"bot_{bot_id}"
    try:
        collections = client.get_collections().collections
        if not any(c.name == collection_name for c in collections):
            return compute_facets([])
    except Exception:
        return compute_facets([])

    payloads = []
    offset = None
    while True:
        response = client.scroll(
            collection_name=collection_name,
            limit=1000,
            offset=offset,
            with_payload=True,
            with_vectors=False,
        )
        points, offset = response
        for p in points:
            payloads.append(p.payload or {})
        if offset is None:
            break

    return compute_facets(payloads)


def get_chunk_counts_by_bot(bot_id: int) -> Dict[int, int]:
    """Count indexed chunks per document from Qdrant (source of truth)."""
    client = get_qdrant_client()
    collection_name = f"bot_{bot_id}"
    counts: Dict[int, int] = {}
    try:
        collections = client.get_collections().collections
        if not any(c.name == collection_name for c in collections):
            return counts
    except Exception:
        return counts

    offset = None
    while True:
        try:
            response = client.scroll(
                collection_name=collection_name,
                limit=1000,
                offset=offset,
                with_payload=True,
                with_vectors=False,
            )
        except Exception:
            break
        points, offset = response
        for p in points:
            meta = (p.payload or {}).get("metadata") or {}
            doc_id = meta.get("doc_id")
            if doc_id is not None:
                counts[doc_id] = counts.get(doc_id, 0) + 1
        if offset is None:
            break
    return counts


def count_document_chunks(bot_id: int, doc_id: int) -> int:
    return get_chunk_counts_by_bot(bot_id).get(doc_id, 0)


def get_collection_stats(bot_id: int) -> Dict[str, Any]:
    client = get_qdrant_client()
    collection_name = f"bot_{bot_id}"
    try:
        info = client.get_collection(collection_name)
        return {
            "collection_name": collection_name,
            "points_count": info.points_count,
            "status": str(info.status),
        }
    except Exception:
        return {"collection_name": collection_name, "points_count": 0, "status": "missing"}


def reindex_all_documents(
    bot_id: int,
    documents: List[Dict[str, Any]],
    embeddings_model,
    chunk_size: int,
    chunk_overlap: int,
) -> Dict[str, Any]:
    delete_bot_collection(bot_id)
    total_chunks = 0
    per_doc: Dict[int, int] = {}
    for doc in documents:
        source_type = doc.get("source_type") or _source_type_from_filename(doc["name"])
        doc_meta = {
            "source_type": source_type,
            "tags": doc.get("tags") or [],
            "source_url": doc.get("source_url") or "",
            "author": doc.get("author") or "",
            "uploaded_at": doc.get("uploaded_at") or "",
            "custom_fields": (doc.get("metadata") or {}).get("custom_fields", {}),
        }
        count = index_document(
            bot_id=bot_id,
            doc_id=doc["id"],
            text=doc["content"],
            name=doc["name"],
            embeddings_model=embeddings_model,
            chunk_size=chunk_size,
            chunk_overlap=chunk_overlap,
            doc_metadata=doc_meta,
        )
        per_doc[doc["id"]] = count
        total_chunks += count
    return {
        "documents_reindexed": len(documents),
        "total_chunks": total_chunks,
        "chunks_per_document": per_doc,
    }
